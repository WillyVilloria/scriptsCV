from docx import Document   # python-docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime 
from googletrans import Translator
from docx.shared import Cm

# Crear documento

class Documento:
    def __init__(self):
        self.doc = Document()
        self.fecha = datetime.now()

    def ecabezado(self):
        # margenes
        section = self.doc.sections[0]
        section.header_distance = Cm(1)
        self.__margenes(section)
        # Acceder al encabezado de la primera sección
        header = self.doc.sections[0].header
        # Limpiar contenido anterior (opcional)
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_run = header_para.add_run("")
        header_run.font.size = Pt(10)         # tamaño de fuente
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 51, 102)   

    def __margenes(self, section):
        # Ajustar márgenes
        section.top_margin = Cm(2.0)       # margen superior
        section.bottom_margin = Cm(2.0)    # margen inferior
        section.left_margin = Cm(2)        # margen izquierdo
        section.right_margin = Cm(2)       # margen derecho

    def pie(self):
        # Acceder al pie de página de la primera sección
        footer = self.doc.sections[0].footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_run = footer_para.add_run("CV Miguel A. Lorenzo Villoria \t")
        footer_run.font.size = Pt(8)         # tamaño de fuente
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 255)
        #Cambiar altura del pie de página
        section = self.doc.sections[0]
        section.footer_distance = Cm(1)

        self.__numero_pagina(footer_para)

    def __numero_pagina(self, footer_para):
        # Insertar "Página X de Y"
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT 
        footer_para.add_run("Página ")
        # Campo PAGE
        run = footer_para.add_run()
        run.font.size = Pt(10)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        footer_para.add_run(" de ")

        # Campo NUMPAGES
        run2 = footer_para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.text = "NUMPAGES"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)

    def estilo(self):
        # Acceder al estilo Normal
        normal_style = self.doc.styles['Normal']

        # Cambiar fuente
        font = normal_style.font
        font.name = 'Cambria'
        font.size = Pt(11)
        font.bold = False
        font.color.rgb = RGBColor(0, 0, 0)  # negro

        # Cambiar párrafo base (alineación, espaciado, etc.)
        paragraph_format = normal_style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.space_before = Pt(6)
        paragraph_format.line_spacing = 1.15

    # Función para añadir un título con color y estilo
    def add_colored_heading(self, text, level=1, color=RGBColor(0, 51, 102)):
        # Crear un heading vacío → esto devuelve un Paragraph
        paragraph = self.doc.add_heading("", level=level)

        # Acceder al estilo correspondiente
        style_name = f"Heading {level}"
        heading_style = self.doc.styles[style_name]

        # Modificar formato de párrafo del estilo (afecta a todos los títulos de ese nivel)
        para_format = heading_style.paragraph_format
        para_format.space_before = Pt(16)   # Espaciado anterior
        para_format.space_after = Pt(8)    # Espaciado posterior

        # Agregar el texto como Run dentro del párrafo
        run = paragraph.add_run(text)
        run.font.color.rgb = color
        run.font.bold = True

        return paragraph

    def cabecera(self):
        # ===== CABECERA =====
        name = self.doc.add_heading("Miguel Ángel Lorenzo Villoria", 0)
        name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        contact = self.doc.add_paragraph(
            "Gijón, Asturias   |   miguelcarter@gmail.com   |    635563966   |   Nacionalidad: Española"
        )
        contact.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        #self.doc.add_paragraph("")

    def Perfil_prof(self):
        translate = Translator()
        # ===== PERFIL PROFESIONAL =====
        self.add_colored_heading("Perfil Profesional", 1)
        perfil1 = self.doc.add_paragraph(
            "Desarrollador / Analista de Datos con formación avanzada en Big Data, Ciencia de Datos y Cloud Computing. Especialista en análisis " \
            "de datos con Python y SQL no SQL, desarrollo de aplicaciones, machine learning y visualización de información."
        )
        perfil1 = self.doc.add_paragraph(
            translate.translate("Desarrollador / Analista de Datos con formación avanzada en Big Data, Ciencia de Datos y Cloud Computing. " \
            "Especialista en análisis de datos con Python y SQL no SQL, desarrollo de aplicaciones, machine learning y visualización de " \
            "información.", dest='es').text
        )
        perfil1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        perfil2 = self.doc.add_paragraph(
            "Experiencia en la implementación de modelos de Machine Learning avanzados, incluyendo redes neuronales profundas, convolucionales (CNN), " \
            "recurrentes (LSTM) y modelos de lenguaje (LLM). Capacidad demostrada en la creación de modelos propios y en la adaptación de " \
            "arquitecturas existentes mediante técnicas de fine-tuning, orientadas a resolver necesidades específicas de negocio y optimizar la " \
            "toma de decisiones."
         )
        perfil2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        perfil3 = self.doc.add_paragraph(
            "Con visión estratégica, aporto la capacidad de transformar datos en conocimiento accionable y de impulsar la innovación tecnológica en " \
            "sectores como banca, consultoría y corporativo."
        )
        perfil3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        perfil4 = self.doc.add_paragraph(
            "Amplios conocimientos de Hadoop y Spark para entornos Big Data."
            
        )
        perfil4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        perfil5 = self.doc.add_paragraph(
            
            "Más de 10 años de experiencia en proyectos tecnológicos y energéticos."
        )
        perfil5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def experiencia_prof(self):
        # ===== EXPERIENCIA PROFESIONAL =====
        self.add_colored_heading("Experiencia Profesional", 1)

        exp1 = self.doc.add_paragraph()
        exp1.add_run("Desarrollador / Analista - Grupo Tecnológico (2022 - Actualidad)").bold = True
        experiencia = [
            "Desarrollo de soluciones de análisis de datos en Python y SQL.",
            "Desarrollo de aplicaciones java (spring boot) para web-Backend.",
            "Automatización de procesos ETL y explotación de bases de datos.",
            "Visualización de resultados y KPIs con Power BI, Qlickview y Tableau.",
            "Desarrollo de proyectos de visión artificial con modelos de segmentación y detección.",
            "Minería de datos, con aplicación de técnicas de clustering y segmentación.",
            "Web scraping, orientado a la extracción, estructuración y análisis de información desde diversas fuentes y análisis con modelos de lenguaje LLM tipo Bert, gpt etc.",
            "Trabajo en entornos cloud y metodologías ágiles."
        ]
        for item in experiencia:
            bullet = self.doc.add_paragraph(f"•\t{item}")
            bullet.paragraph_format.left_indent = Pt(18)         # sangría para alinear con el texto
            bullet.paragraph_format.first_line_indent = Pt(-18)  # primera línea “sale” la viñeta
            bullet.paragraph_format.space_after = Pt(0)
            bullet.paragraph_format.line_spacing = 1.0 

        exp2 = self.doc.add_paragraph()
        exp2.add_run("Jefe de Operaciones / Servicio - UTE ESE CLECE - Gas Natural Fenosa (2012 - 2022)").bold = True
        operaciones = [
            "Dirección de proyectos de eficiencia energética de gran escala (>10MW).",
            "Desarrollo de aplicación en C# con MySQL para gestión de datos operativos.",
            "Coordinación de equipos multidisciplinares y control de KPIs.",
            "Implementación de soluciones de monitorización digital de instalaciones."
        ]
        for item in operaciones:
            bullet = self.doc.add_paragraph(f"•\t{item}")
            bullet.paragraph_format.left_indent = Pt(18)         # sangría para alinear con el texto
            bullet.paragraph_format.first_line_indent = Pt(-18)  # primera línea “sale” la viñeta
            bullet.paragraph_format.space_after = Pt(0)
            bullet.paragraph_format.line_spacing = 1.0 

        exp3 = self.doc.add_paragraph()
        exp3.add_run("Responsable de Proyectos e Instalaciones - Energía y Servicios Energéticos Plus S.A.\n").bold = True
        exp3.add_run("Gestión de proyectos energéticos con preferencia en control de datos técnicos y procesos.")

        exp4 = self.doc.add_paragraph()
        exp4.add_run("Director Técnico - Audiner Asesores Energéticos S.L.").bold = True

        exp5 = self.doc.add_paragraph()
        exp5.add_run("Jefe de Sector - Giroa S.A.U. (Grupo Veolia), Asturias").bold = True

        exp6 = self.doc.add_paragraph()
        exp6.add_run("Jefe de Servicio - Integra - Clece S.A. (Grupo ACS - Dragados)").bold = True

        exp7 = self.doc.add_paragraph()
        exp7.add_run("Inspector Técnico - ECA S.A.").bold = True

    def logros(self):
        # ===== LOGROS DESTACADOS =====
        self.add_colored_heading("Logros Destacados", 1)
        logros = [
            "Estudio analítico para el aumento de potencia instalada en planta de producción de frío industrial en MercaMadrid.",
            "Optimización del consumo energético en Complejo Ministerial mediante análisis predictivo y modelado estadístico.",
            "Proyectos de visión artificial, aplicando redes neuronales profundas para la detección y clasificación de patrones en imágenes.",
            "Desarrollo de aplicaciones web backend en Java, integradas con bases de datos y orientadas a la gestión y explotación de información.",
            "Minería de datos y clustering para segmentación y descubrimiento de insights en grandes volúmenes de información."]
        
        for item in logros:
            bullet = self.doc.add_paragraph(f"•\t{item}")
            bullet.paragraph_format.left_indent = Pt(20)         # sangría para alinear con el texto
            bullet.paragraph_format.first_line_indent = Pt(-18)  # primera línea “sale” la viñeta
            bullet.paragraph_format.space_after = Pt(0)
            bullet.paragraph_format.line_spacing = 1.0 

    def formacion(self):
        # ===== FORMACIÓN =====
        self.add_colored_heading("Formación Académica", 1)
        formacion = [
            "Máster Universitario en Big Data y Ciencia de Datos - Universidad Internacional de Valencia (VIU)",
            "Experto Universitario en Programación Python y Ciencia de Datos - VIU",    
            "Experto Universitario en DevOps y Cloud Computing - UNIR",
            "Ingeniero Técnico de Minas - Universidad de Oviedo"
        ]
        for item in formacion:
            bullet = self.doc.add_paragraph()
            run = bullet.add_run(f"•\t{item}")
            run.bold = True
            bullet.paragraph_format.left_indent = Pt(20)         # sangría para alinear con el texto
            bullet.paragraph_format.first_line_indent = Pt(-18)  # primera línea “sale” la viñeta
            bullet.paragraph_format.space_after = Pt(0)
            bullet.paragraph_format.line_spacing = 1.0 
        otros_cursos = "Otros cursos: Máster en Gestión de Calidad y Medio Ambiente, Técnico de Prevención de Riesgos Laborales, Autocad, Soldadura técnica."
        self.doc.add_paragraph(otros_cursos)

    def habilidades(self):
        # ===== HABILIDADES TÉCNICAS =====
        self.add_colored_heading("Habilidades Técnicas", 1)
        habilidades = [
            "Python (pandas, NumPy, scikit-learn, matplotlib, seaborn), SQL, C# .NET, Java, C/C++.",
            "Machine Learning, Big Data (Hadoop, Spark).",
            "Cloud Computing (AWS, Azure, GCP), DevOps, Git.",
            "Power BI, Tableau, Qlikview.",
            "Metodologías Agile/Scrum."
        ]
        for item in habilidades:
            bullet = self.doc.add_paragraph(f"•\t{item}")
            bullet.paragraph_format.left_indent = Pt(20)         # sangría para alinear con el texto
            bullet.paragraph_format.first_line_indent = Pt(-18)  # primera línea “sale” la viñeta
            bullet.paragraph_format.space_after = Pt(0)
            bullet.paragraph_format.line_spacing = 1.0 

    def idiomas(self):
        # ===== IDIOMAS =====
        self.add_colored_heading("Idiomas", 1)
        self.doc.add_paragraph("• Español (Nativo)\n• Inglés (Nivel alto – lectura, escritura y expresión oral)")

    def guardar(self):
        # Guardar documento
        output_path_visual = f"../CV/LCV_Miguel_Angel_Lorenzo_Villoria_Data_Visual{self.fecha.day}.docx"
        self.doc.save(output_path_visual)

        output_path_visual

if __name__ == "__main__":
    document = Documento()
    document.ecabezado()
    document.pie()
    document.estilo()
    document.cabecera()
    document.Perfil_prof()
    document.experiencia_prof()
    document.logros()
    document.formacion()
    document.habilidades()
    document.idiomas()
    document.guardar()